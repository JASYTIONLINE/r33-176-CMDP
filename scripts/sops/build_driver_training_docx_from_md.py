"""Build Driver Training SOP DOCX from source/sops/10.10.1-sop-driver-training.md (source of truth, AR 25-50).

Output (canonical shipped copy):
  assets/word/sops/10.10.1-sop-driver-training.docx

Requires: pip install python-docx
"""
from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

from scripts.paths import repo_root
from scripts.word.convert_sops_to_docx import (  # noqa: E402
    convert_markdown_table_to_word,
    create_paragraph_with_formatting,
    setup_arial_font,
)

_REPO = repo_root()

MD_PATH = _REPO / "source/sops/10.10.1-sop-driver-training.md"
OUT_DOCX = _REPO / "assets/word/sops/10.10.1-sop-driver-training.docx"

SECTION_LIKE = re.compile(r"^##\s+(\d+-\d+)\.\s+")
CHAPTER_LINE = re.compile(r"^#\s+Chapter\s+", re.I)
ENUM_CHAPTER = re.compile(r"^#\s+\d+\.\s+[A-Z]")
LIST_INTRO = re.compile(
    r"(consist of|framework consisting of|agencies including|relationships including|designated organizations? and agencies including):\s*$",
    re.I,
)


def _introduces_plain_list(prev_line: str) -> bool:
    t = prev_line.strip()
    if LIST_INTRO.search(t):
        return True
    if t.endswith(":") and len(t) < 180:
        tl = t.lower()
        if tl.endswith("including:") or tl.endswith("consist of:"):
            return True
    return False


SENTENCE_START = re.compile(
    r"^(The |All |Units |Battalions |Companies |Commanders |Driver training |This SOP |Subordinate |The Battalion|The Brigade|When |If |Operators |Personnel )\w",
)


def _margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.line_spacing = 1.0


def normalize_text(s: str) -> str:
    s = html.unescape(s)
    s = re.sub(r"\\([\[\]])", r"\1", s)
    return s


def load_normalized_lines(path: Path) -> list[str]:
    raw = path.read_text(encoding="utf-8")
    return [normalize_text(line) for line in raw.splitlines()]


def should_skip_nav(line: str) -> bool:
    t = line.strip()
    return ("[[" in t and "]]" in t) or ("sops.index" in t and "|" in t)


def preprocess_plain_list_blocks(lines: list[str]) -> list[str]:
    """Lines after list-intros become '- ' bullets (in-memory only)."""
    out: list[str] = []
    i = 0
    n = len(lines)
    while i < n:
        ln = lines[i]
        if _introduces_plain_list(ln):
            out.append(ln)
            i += 1
            while i < n and not lines[i].strip():
                out.append(lines[i])
                i += 1
            while i < n:
                nxt = lines[i]
                s = nxt.strip()
                if not s:
                    out.append(nxt)
                    i += 1
                    break
                if s.startswith("#") or s.startswith("##") or s.startswith("- "):
                    break
                if s.startswith("|"):
                    break
                if len(s) > 160:
                    break
                if SENTENCE_START.match(s) and len(s) > 55:
                    break
                out.append(f"- {s}")
                i += 1
            continue
        out.append(ln)
        i += 1
    return out


def add_para_runs_bold(doc: Document, text: str, **kwargs) -> None:
    """Paragraph with **bold** preserved; AR spacing via create_paragraph_with_formatting base."""
    alignment = kwargs.get("alignment")
    indent_left = kwargs.get("indent_left", 0)
    first_line_indent = kwargs.get("first_line_indent", 0)
    is_heading = kwargs.get("is_heading", False)
    default_bold = kwargs.get("bold", False)

    paragraph = doc.add_paragraph()
    if alignment is not None:
        paragraph.alignment = alignment
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0) if is_heading else Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    if indent_left:
        pf.left_indent = Inches(indent_left)
    if first_line_indent:
        pf.first_line_indent = Inches(first_line_indent)

    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            setup_arial_font(run)
        else:
            run = paragraph.add_run(part)
            if default_bold:
                run.bold = True
            setup_arial_font(run)


def add_bullet(doc: Document, text: str, indent_extra: float = 0.0) -> None:
    paragraph = doc.add_paragraph()
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    pf.left_indent = Inches(0.5 + indent_extra)
    pf.first_line_indent = Inches(-0.22)
    parts = re.split(r"(\*\*[^*]+\*\*)", text.strip())
    bullet = paragraph.add_run("• ")
    setup_arial_font(bullet)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            setup_arial_font(run)
        else:
            run = paragraph.add_run(part)
            setup_arial_font(run)


def find_body_start(lines: list[str]) -> int:
    """Maint-style SOPs: body begins at ## 1. PURPOSE."""
    for i, ln in enumerate(lines):
        if should_skip_nav(ln):
            continue
        if re.match(r"^##\s+1\.\s+PURPOSE\b", ln.strip(), re.I):
            return i
    return -1


def find_summary_index(lines: list[str]) -> int:
    """Legacy driver SOPs: body began at # SUMMARY of CHANGE."""
    for i, ln in enumerate(lines):
        if should_skip_nav(ln):
            continue
        if re.match(r"^#\s+SUMMARY\b", ln.strip(), re.I):
            return i
    return -1


def build_cover(doc: Document, lines: list[str], end: int) -> None:
    """Cover through line end (exclusive of first body heading e.g. ## 1. PURPOSE)."""
    i = 0
    addr_buf: list[str] = []

    def flush_addr() -> None:
        nonlocal addr_buf
        if not addr_buf:
            return
        bucket: list[str] = []
        paras: list[str] = []
        for s in addr_buf:
            if s.strip():
                bucket.append(s.strip())
            else:
                if bucket:
                    paras.append(" ".join(bucket))
                    bucket = []
        if bucket:
            paras.append(" ".join(bucket))
        for p in paras:
            if p.strip():
                create_paragraph_with_formatting(doc, p.strip(), bold=False)
        addr_buf = []

    while i < end:
        raw = lines[i]
        if should_skip_nav(raw):
            i += 1
            continue
        ts = raw.strip()
        if not ts:
            # Keep blank lines while accumulating letterhead so flush_addr splits paragraphs
            # (same effective layout as the maint SOP Word output).
            if addr_buf:
                addr_buf.append(raw)
            i += 1
            continue

        if ts == "---":
            i += 1
            continue

        m_unit = re.fullmatch(r"\*\*(.+)\*\*", ts)
        if m_unit and "|" not in ts:
            flush_addr()
            create_paragraph_with_formatting(
                doc,
                m_unit.group(1).strip(),
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                is_heading=True,
            )
            i += 1
            continue

        if (
            ts.startswith("## ")
            and ts[3:].strip().lower().startswith("administrative data")
        ):
            flush_addr()
            i += 1
            table_lines: list[str] = []
            while i < end:
                u = lines[i].strip()
                if not u:
                    i += 1
                    continue
                if u.startswith("## ") or u.startswith("# "):
                    break
                if u.startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                    continue
                i += 1
            if table_lines:
                convert_markdown_table_to_word(doc, table_lines)
                create_paragraph_with_formatting(doc, "", is_empty=True)
            continue

        if ts.startswith("# ") and "STANDARD OPERATING PROCEDURES" in ts.upper():
            flush_addr()
            create_paragraph_with_formatting(
                doc,
                ts[2:].strip(),
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                is_heading=True,
            )
            i += 1
            while i < end and not lines[i].strip():
                i += 1
            if i < end and lines[i].strip().startswith("## "):
                short_title = lines[i].strip()[3:].strip()
                by_line = ""
                i += 1
                while i < end and not lines[i].strip():
                    i += 1
                if i < end and lines[i].strip().upper().startswith("BY ORDER"):
                    by_line = lines[i].strip()
                    i += 1
                create_paragraph_with_formatting(
                    doc,
                    short_title,
                    bold=True,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    is_heading=True,
                )
                if by_line:
                    create_paragraph_with_formatting(
                        doc, by_line, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER
                    )
            continue

        if ts.startswith("# ") and ts.upper().strip().startswith("# DRAFT"):
            flush_addr()
            create_paragraph_with_formatting(
                doc, ts[2:].strip(), bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT, is_heading=True
            )
            i += 1
            continue

        if ts.startswith("## ") and ts[3:].strip().upper().startswith("OFFICIAL"):
            flush_addr()
            i += 1
            blk: list[str] = []
            while i < end:
                u = lines[i].strip()
                if not u:
                    i += 1
                    continue
                if u.startswith("## ") or u.startswith("# "):
                    break
                blk.append(u)
                i += 1
            official_txt = "\n".join(blk)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r0 = p.add_run("OFFICIAL:")
            r0.bold = True
            setup_arial_font(r0)
            for ln in blk:
                br = p.add_run()
                br.add_break()
                setup_arial_font(br)
                rr = p.add_run(ln)
                setup_arial_font(rr)
            p.paragraph_format.space_after = Pt(12)
            continue

        if ts.startswith("## "):
            flush_addr()
            label = ts[3:].strip().rstrip(":")
            i += 1
            body_lines: list[str] = []
            while i < end:
                u = lines[i].strip()
                if not u:
                    i += 1
                    continue
                if u.startswith("## ") or u.startswith("# "):
                    break
                body_lines.append(u)
                i += 1
            label_clean = label.rstrip(".")
            if label.lower().startswith("distribution"):
                create_paragraph_with_formatting(doc, f"{label_clean}:", bold=True)
                for bl in body_lines:
                    if bl.strip().upper() == "A":
                        create_paragraph_with_formatting(
                            doc, "A", bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER
                        )
                    else:
                        create_paragraph_with_formatting(doc, bl.strip())
            else:
                body_txt = " ".join(body_lines).strip()
                if body_txt:
                    add_para_runs_bold(
                        doc,
                        f"**{label_clean}:** {body_txt}",
                    )
                else:
                    create_paragraph_with_formatting(
                        doc, f"{label_clean}:", bold=True, is_heading=True
                    )
            continue

        addr_buf.append(raw)
        i += 1

    flush_addr()
    create_paragraph_with_formatting(doc, "", is_empty=True)


def build_body(doc: Document, lines: list[str], start: int) -> None:
    i = start
    n = len(lines)
    first_after_chapter = False

    while i < n:
        raw = lines[i]
        if should_skip_nav(raw):
            i += 1
            continue
        lstripped = raw.lstrip()
        lead_spaces = len(raw) - len(lstripped)
        line = lstripped.strip()
        if not line:
            i += 1
            continue

        if line == "---":
            i += 1
            continue

        if line.startswith("|"):
            table_lines: list[str] = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            convert_markdown_table_to_word(doc, table_lines)
            create_paragraph_with_formatting(doc, "", is_empty=True)
            continue

        if line.startswith("# ") and not line.startswith("##"):
            inner = line[2:].strip()
            u = inner.upper()
            if u.startswith("SUMMARY"):
                create_paragraph_with_formatting(
                    doc, inner, bold=True, is_heading=True
                )
                first_after_chapter = False
            elif u.startswith("CHAPTER ") or ENUM_CHAPTER.match(line):
                create_paragraph_with_formatting(
                    doc, inner, bold=True, is_heading=True
                )
                first_after_chapter = True
            else:
                create_paragraph_with_formatting(
                    doc, inner, bold=True, is_heading=True
                )
                first_after_chapter = False
            i += 1
            continue

        if line.startswith("## ") and not line.startswith("###"):
            subt = line[3:].strip()
            if SECTION_LIKE.match(line):
                if first_after_chapter:
                    create_paragraph_with_formatting(
                        doc, subt, bold=True, is_heading=True
                    )
                    first_after_chapter = False
                else:
                    create_paragraph_with_formatting(
                        doc, subt, bold=True, indent_left=0.12, is_heading=True
                    )
            else:
                create_paragraph_with_formatting(
                    doc, subt, bold=True, is_heading=True
                )
                first_after_chapter = False
            i += 1
            continue

        if line.startswith("### "):
            create_paragraph_with_formatting(
                doc,
                line[4:].strip(),
                bold=True,
                indent_left=0.18,
                is_heading=True,
            )
            i += 1
            continue

        if line.startswith("!["):
            i += 1
            continue

        if line.startswith("**APPROVAL:**"):
            for _ in range(4):
                create_paragraph_with_formatting(doc, "", is_empty=True)
            create_paragraph_with_formatting(doc, "APPROVAL:", bold=True)
            create_paragraph_with_formatting(doc, "", is_empty=True)
            i += 1
            continue

        if line.startswith("**[LAST NAME") or line.startswith("[LAST NAME"):
            sig_text = line.replace("**", "").replace("[", "").replace("]", "").strip()
            create_paragraph_with_formatting(doc, sig_text, bold=True)
            i += 1
            if i < n:
                next_line = (
                    lines[i].strip().replace("[", "").replace("]", "").strip()
                )
                if (
                    next_line
                    and not next_line.startswith("![")
                    and not next_line.startswith("**CMDP")
                ):
                    create_paragraph_with_formatting(doc, next_line)
                    i += 1
            if i < n:
                next_line = lines[i].strip()
                if next_line == "Commanding":
                    create_paragraph_with_formatting(
                        doc,
                        next_line,
                        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                    )
                    i += 1
            continue

        if line.startswith("**[") and line.endswith("**") and "," in line:
            sig_text = line.replace("**", "").replace("[", "").replace("]", "").strip()
            create_paragraph_with_formatting(doc, sig_text, bold=True)
            i += 1
            if i < n:
                next_line = (
                    lines[i].strip().replace("[", "").replace("]", "").strip()
                )
                if (
                    next_line
                    and not next_line.startswith("![")
                    and not next_line.startswith("**CMDP")
                ):
                    create_paragraph_with_formatting(doc, next_line)
                    i += 1
            if i < n:
                next_line = lines[i].strip()
                if next_line == "Commanding":
                    create_paragraph_with_formatting(
                        doc,
                        next_line,
                        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                    )
                    i += 1
            continue

        if line.startswith("**CMDP Reference:**") or line.startswith(
            "**CMDP Reference"
        ):
            ref_text = line.replace("**", "").strip()
            create_paragraph_with_formatting(doc, ref_text, bold=True)
            i += 1
            continue

        if lstripped.startswith("- "):
            extra = 0.28 if lead_spaces >= 2 else 0.0
            add_bullet(doc, lstripped[2:].strip(), indent_extra=extra)
            i += 1
            continue

        bucket = [line]
        i += 1
        while i < n:
            nxt = lines[i]
            if should_skip_nav(nxt):
                i += 1
                continue
            nx = nxt.strip()
            if not nx:
                break
            if nx == "---":
                break
            if nx.startswith("#"):
                break
            if nx.startswith("## ") or nx.startswith("### "):
                break
            if nx.startswith("|"):
                break
            if nxt.lstrip().startswith("- "):
                break
            bucket.append(nx)
            i += 1
        merged = " ".join(bucket)
        add_para_runs_bold(doc, merged)


def build() -> None:
    if not MD_PATH.is_file():
        raise SystemExit(f"Missing {MD_PATH}")

    lines = load_normalized_lines(MD_PATH)
    lines = preprocess_plain_list_blocks(lines)
    body_start = find_body_start(lines)
    if body_start < 0:
        body_start = find_summary_index(lines)
    if body_start < 0:
        raise SystemExit(
            "Expected '## 1. PURPOSE' (maint-style) or '# SUMMARY of CHANGE' (legacy)."
        )

    doc = Document()
    _margins(doc)

    build_cover(doc, lines[:body_start], len(lines[:body_start]))
    build_body(doc, lines, body_start)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

    def save_doc(path: Path) -> None:
        doc.save(path)

    try:
        save_doc(OUT_DOCX)
        print(f"Wrote {OUT_DOCX.relative_to(_REPO)}")
    except PermissionError:
        fallback = _REPO / "drafts" / OUT_DOCX.name
        fallback.parent.mkdir(parents=True, exist_ok=True)
        save_doc(fallback)
        print(
            f"WARNING: Could not write {OUT_DOCX.name} (file may be open in Word). "
            f"Wrote {fallback.relative_to(_REPO)} instead; close the file and re-run."
        )


if __name__ == "__main__":
    build()
