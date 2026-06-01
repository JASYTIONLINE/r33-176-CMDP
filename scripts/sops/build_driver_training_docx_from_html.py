"""Build Driver Training SOP DOCX from sops/10.10.1-sop-driver-training.html.

For authoritative Markdown source and AR 25-50 layout, use
``build_driver_training_docx_from_md.py`` instead.

Outputs:
  assets/word/sops/10.10.1-sop-driver-training.docx

Requires: pip install python-docx beautifulsoup4
"""
from __future__ import annotations

from pathlib import Path

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

from scripts.paths import repo_root, sop_publish_dir
from scripts.word.convert_sops_to_docx import create_paragraph_with_formatting, setup_arial_font

_REPO = repo_root()
HTML_PATH = sop_publish_dir() / "10.10.1-sop-driver-training.html"
OUT_DOCX = _REPO / "assets/word/sops/10.10.1-sop-driver-training.docx"


def _margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.line_spacing = 1.0


def add_html_table(doc: Document, table) -> None:
    rows_data: list[list[str]] = []
    for tr in table.find_all("tr"):
        cells = tr.find_all(["th", "td"])
        rows_data.append([c.get_text(" ", strip=True) for c in cells])
    if not rows_data:
        return
    ncols = max(len(r) for r in rows_data)
    t = doc.add_table(rows=len(rows_data), cols=ncols)
    t.style = "Light Grid Accent 1"
    for i, row_data in enumerate(rows_data):
        for j in range(ncols):
            cell_text = row_data[j] if j < len(row_data) else ""
            cell = t.rows[i].cells[j]
            cell.paragraphs[0].clear()
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            setup_arial_font(run)
            if i == 0:
                run.bold = True


def _alignment_from_cover_classes(classes: list[str]) -> WD_ALIGN_PARAGRAPH | None:
    cs = frozenset(classes or [])
    if "cover-center" in cs:
        return WD_ALIGN_PARAGRAPH.CENTER
    if "cover-right" in cs:
        return WD_ALIGN_PARAGRAPH.RIGHT
    if "cover-left" in cs:
        return WD_ALIGN_PARAGRAPH.LEFT
    return None


def _add_inline_with_breaks(paragraph, node) -> None:
    from docx.enum.text import WD_BREAK

    for child in node.children:
        if isinstance(child, NavigableString):
            t = str(child)
            if not t.strip():
                continue
            run = paragraph.add_run(t)
            setup_arial_font(run)
        elif getattr(child, "name", None) == "br":
            paragraph.add_run().add_break(WD_BREAK.LINE)
        elif getattr(child, "name", None) in ("strong", "b"):
            run = paragraph.add_run(child.get_text())
            run.bold = True
            setup_arial_font(run)
        elif getattr(child, "name", None) == "a":
            href = child.get("href", "").strip()
            run = paragraph.add_run(child.get_text() + (f" ({href})" if href else ""))
            setup_arial_font(run)
        elif getattr(child, "name", None) is not None:
            _add_inline_with_breaks(paragraph, child)


def add_paragraph_from_p(doc: Document, p_tag, alignment: WD_ALIGN_PARAGRAPH | None = None) -> None:
    if not p_tag.get_text(strip=True):
        return
    paragraph = doc.add_paragraph()
    if alignment is not None:
        paragraph.alignment = alignment
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    _add_inline_with_breaks(paragraph, p_tag)


def add_ul(doc: Document, ul, indent_level: float = 0.0) -> None:
    base_left = 0.5 + indent_level
    for li in ul.find_all("li", recursive=False):
        nested_lists: list = []
        bucket: list = []
        for ch in li.children:
            if getattr(ch, "name", None) in ("ul", "ol"):
                nested_lists.append(ch)
            else:
                bucket.append(ch)

        texty = False
        for node in bucket:
            if isinstance(node, NavigableString) and str(node).strip():
                texty = True
                break
            if getattr(node, "name", None) is not None:
                texty = True
                break

        if texty:
            para = doc.add_paragraph()
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(12)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.left_indent = Inches(base_left)
            pf.first_line_indent = Inches(-0.25)
            pr = para.add_run("• ")
            setup_arial_font(pr)
            for node in bucket:
                if isinstance(node, NavigableString):
                    t = str(node)
                    if t.strip():
                        r = para.add_run(t)
                        setup_arial_font(r)
                elif getattr(node, "name", None) in ("strong", "b"):
                    r = para.add_run(node.get_text())
                    r.bold = True
                    setup_arial_font(r)
                elif getattr(node, "name", None) == "a":
                    href = node.get("href", "").strip()
                    r = para.add_run(node.get_text() + (f" ({href})" if href else ""))
                    setup_arial_font(r)
                elif getattr(node, "name", None):
                    _add_inline_with_breaks(para, node)

        for nl in nested_lists:
            if nl.name == "ul":
                add_ul(doc, nl, indent_level + 0.25)
            else:
                add_ol(doc, nl, indent_level + 0.25)


def add_ol(doc: Document, ol, indent_level: float = 0.0) -> None:
    base_left = 0.5 + indent_level
    n = 1
    for li in ol.find_all("li", recursive=False):
        nested_lists: list = []
        bucket: list = []
        for ch in li.children:
            if getattr(ch, "name", None) in ("ul", "ol"):
                nested_lists.append(ch)
            else:
                bucket.append(ch)

        texty = False
        for node in bucket:
            if isinstance(node, NavigableString) and str(node).strip():
                texty = True
                break
            if getattr(node, "name", None) is not None:
                texty = True
                break

        item_num = n
        n += 1

        if texty:
            para = doc.add_paragraph()
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(12)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.left_indent = Inches(base_left)
            pf.first_line_indent = Inches(-0.25)
            pr = para.add_run(f"{item_num}. ")
            setup_arial_font(pr)
            for node in bucket:
                if isinstance(node, NavigableString):
                    t = str(node)
                    if t.strip():
                        r = para.add_run(t)
                        setup_arial_font(r)
                elif getattr(node, "name", None) in ("strong", "b"):
                    r = para.add_run(node.get_text())
                    r.bold = True
                    setup_arial_font(r)
                elif getattr(node, "name", None) == "a":
                    href = node.get("href", "").strip()
                    r = para.add_run(node.get_text() + (f" ({href})" if href else ""))
                    setup_arial_font(r)
                elif getattr(node, "name", None):
                    _add_inline_with_breaks(para, node)

        for nl in nested_lists:
            if nl.name == "ul":
                add_ul(doc, nl, indent_level + 0.25)
            else:
                add_ol(doc, nl, indent_level + 0.25)


def _letterhead_cover(doc: Document, h1_text: str) -> None:
    for line in (
        "TEXAS MILITARY DEPARTMENT",
        "TEXAS ARMY NATIONAL GUARD (TXARNG)",
        "P.O. BOX 5218",
        "AUSTIN, TX 78763-5218",
        "[INSERT DATE]",
    ):
        create_paragraph_with_formatting(doc, line, bold=False, alignment=None)
    create_paragraph_with_formatting(doc, "", is_empty=True)
    short = "176TH EN BDE DRIVER TRAINING SOP"
    create_paragraph_with_formatting(
        doc, h1_text, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, is_heading=True
    )
    create_paragraph_with_formatting(
        doc, short, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, is_heading=True
    )
    by_line = "By Order of the Commander, 176th Engineer Brigade(TXARNG)"
    create_paragraph_with_formatting(
        doc, by_line, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    create_paragraph_with_formatting(doc, "", is_empty=True)
    draft = doc.add_paragraph()
    draft.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = draft.add_run("DRAFT-For Reference Only")
    r.bold = True
    setup_arial_font(r)
    create_paragraph_with_formatting(doc, "", is_empty=True)
    create_paragraph_with_formatting(doc, "", is_empty=True)



def build() -> None:
    html = HTML_PATH.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "html.parser")
    article = soup.select_one("article.sop-body")

    doc = Document()
    _margins(doc)

    if article:
        cover = article.select_one("div.cover-block")
        if cover:
            for p in cover.find_all("p", recursive=False):
                if not p.get_text(strip=True):
                    continue
                align = _alignment_from_cover_classes(p.get("class") or [])
                add_paragraph_from_p(doc, p, alignment=align)
        elems: list = []
        for el in article.children:
            nm = getattr(el, "name", None)
            if not nm:
                continue
            if nm == "div" and el.get("class") and "cover-block" in el.get("class"):
                continue
            elems.append(el)
    else:
        main = soup.find("main")
        if not main:
            raise SystemExit("No <main> / article.sop-body in driver training HTML.")
        elems = []
        for el in main.children:
            nm = getattr(el, "name", None)
            if not nm:
                continue
            if el.name == "section" and el.get("id") == "draft-banner":
                continue
            if el.name == "div" and "sop-actions" in " ".join(el.get("class", [])):
                continue
            elems.append(el)
        while elems and elems[0].name == "hr":
            elems.pop(0)
        idx = 0
        while idx < len(elems) and elems[idx].name == "hr":
            idx += 1
        elems = elems[idx:]
        legacy_h1 = ""
        legacy_unit = ""
        if elems and elems[0].name == "h1":
            legacy_h1 = elems[0].get_text(" ", strip=True)
            elems = elems[1:]
        if elems and elems[0].name == "p":
            legacy_unit = elems[0].get_text(" ", strip=True)
            elems = elems[1:]
        while elems and elems[0].name == "hr":
            elems = elems[1:]
        _letterhead_cover(doc, legacy_h1 or "DRIVER TRAINING STANDARD OPERATING PROCEDURE")
        if legacy_unit:
            p_unit = doc.add_paragraph()
            p_unit.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_unit = p_unit.add_run(legacy_unit)
            r_unit.bold = True
            setup_arial_font(r_unit)
            p_unit.paragraph_format.space_after = Pt(12)
            create_paragraph_with_formatting(doc, "", is_empty=True)

    for el in elems:
        if el.name == "hr":
            create_paragraph_with_formatting(doc, "", is_empty=True)
        elif el.name == "h1":
            create_paragraph_with_formatting(
                doc,
                el.get_text(" ", strip=True),
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                is_heading=True,
            )
        elif el.name == "h2":
            cls = " ".join(el.get("class", []))
            if "chapter" in cls.split():
                create_paragraph_with_formatting(
                    doc, el.get_text(" ", strip=True), bold=True, is_heading=True
                )
            else:
                create_paragraph_with_formatting(
                    doc, el.get_text(" ", strip=True), bold=True, is_heading=True
                )
        elif el.name == "h3":
            create_paragraph_with_formatting(
                doc, el.get_text(" ", strip=True), bold=True, indent_left=0.25, is_heading=True
            )
        elif el.name == "h4":
            create_paragraph_with_formatting(
                doc, el.get_text(" ", strip=True), bold=True, indent_left=0.35, is_heading=True
            )
        elif el.name == "p":
            add_paragraph_from_p(doc, el)
        elif el.name == "table":
            add_html_table(doc, el)
            create_paragraph_with_formatting(doc, "", is_empty=True)
        elif el.name == "ul":
            add_ul(doc, el)
        elif el.name == "ol":
            add_ol(doc, el)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX.relative_to(_REPO)}")


if __name__ == "__main__":
    build()
