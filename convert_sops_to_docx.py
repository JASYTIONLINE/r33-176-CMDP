"""
Convert SOP markdown files to Word documents formatted per AR 25-50.
Requires: pip install python-docx
"""
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def setup_arial_font(run):
    """Set font to 12pt Arial for a run"""
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    # Set East Asian font to Arial for compatibility
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Arial')

def create_paragraph_with_formatting(doc, text='', bold=False, alignment=None, 
                                     indent_left=0, first_line_indent=0, 
                                     is_empty=False, is_heading=False):
    """Create a paragraph with AR 25-50 formatting"""
    paragraph = doc.add_paragraph()
    
    if alignment:
        paragraph.alignment = alignment
    
    # Set paragraph formatting per AR 25-50
    para_format = paragraph.paragraph_format
    para_format.space_before = Pt(0)
    # Double spacing between paragraphs = 12pt (one line of 12pt font)
    para_format.space_after = Pt(12) if not is_empty and not is_heading else Pt(0)
    para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para_format.line_spacing = 1.0
    
    if indent_left > 0:
        para_format.left_indent = Inches(indent_left)
    if first_line_indent != 0:
        para_format.first_line_indent = Inches(first_line_indent)
    
    if text:
        # Handle bold text markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                setup_arial_font(run)
            else:
                run = paragraph.add_run(part)
                if bold:
                    run.bold = True
                setup_arial_font(run)
    else:
        # Empty paragraph - still set font
        run = paragraph.add_run()
        setup_arial_font(run)
    
    return paragraph

def convert_markdown_table_to_word(doc, table_lines):
    """Convert markdown table to Word table"""
    # Parse markdown table
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line or not line.startswith('|'):
            continue
        # Remove leading/trailing pipes and split
        cells = [cell.strip() for cell in line.strip('|').split('|')]
        # Skip separator row (contains only dashes and colons)
        if all(cell.replace('-', '').replace(':', '').strip() == '' for cell in cells):
            continue
        if cells:  # Only add non-empty rows
            rows.append(cells)
    
    if not rows:
        return
    
    # Determine number of columns from first row
    num_cols = len(rows[0]) if rows else 2
    
    # Create Word table
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(rows):
        for j in range(num_cols):
            cell_text = row_data[j] if j < len(row_data) else ''
            cell = table.rows[i].cells[j]
            # Clear default paragraph
            cell.paragraphs[0].clear()
            # Process text with bold markers
            para = cell.paragraphs[0]
            parts = re.split(r'(\*\*[^*]+\*\*)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                    setup_arial_font(run)
                else:
                    run = para.add_run(part)
                    setup_arial_font(run)
            # Make first row bold (header row)
            if i == 0:
                for run in para.runs:
                    run.bold = True

def convert_markdown_to_docx(md_file_path, output_dir):
    """Convert a markdown SOP file to DOCX with AR 25-50 formatting"""
    
    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set document margins per AR 25-50: 1 inch on all sides
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set default font to Arial 12pt in Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style.paragraph_format.space_before = Pt(0)
    # AR 25-50: double-space between paragraphs
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.line_spacing = 1.0
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip markdown navigation links
        if line.startswith('[[') and ('|' in line or 'sops.index' in line or 'cmdp-index' in line):
            i += 1
            continue
        
        # Skip image references (signature images)
        if line.startswith('!['):
            i += 1
            continue
        
        # Handle horizontal rules (---) - add spacing
        if line == '---':
            create_paragraph_with_formatting(doc, '', is_empty=True)
            i += 1
            continue
        
        # Handle empty lines
        if not line:
            if not in_table:
                create_paragraph_with_formatting(doc, '', is_empty=True)
            i += 1
            continue
        
        # Handle markdown tables
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(lines[i])  # Use original line with spacing
            i += 1
            # Continue collecting table rows
            continue
        
        # If we were in a table and hit a non-table line, process the table
        if in_table and not line.startswith('|'):
            convert_markdown_table_to_word(doc, table_lines)
            table_lines = []
            in_table = False
            # Don't increment i, process this line normally
        
        # Handle H1 title (centered, bold)
        if line.startswith('# ') and not line.startswith('##'):
            title = line[2:].strip()
            para = create_paragraph_with_formatting(doc, title, bold=True, 
                                                   alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                                   is_heading=True)
            i += 1
            continue
        
        # Handle unit name (bold, centered)
        if line.startswith('**') and line.endswith('**') and '[UNIT NAME]' in line:
            unit_name = line.replace('**', '').strip()
            create_paragraph_with_formatting(doc, unit_name, bold=True,
                                           alignment=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue
        
        # Handle numbered sections (## 1. PURPOSE)
        if re.match(r'^##\s+\d+\.', line):
            section_text = re.sub(r'^##\s+', '', line).strip()
            para = create_paragraph_with_formatting(doc, section_text, bold=True, is_heading=True)
            i += 1
            continue
        
        # Handle subsections (### 4.1 Commander)
        if re.match(r'^###\s+\d+\.\d+', line):
            subsection_text = re.sub(r'^###\s+', '', line).strip()
            # Remove bold markers but keep content
            clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', subsection_text)
            para = create_paragraph_with_formatting(doc, clean_text, bold=True,
                                                   indent_left=0.25, is_heading=True)
            i += 1
            continue
        
        # Handle lettered sub-items (a., b., c.)
        if re.match(r'^[a-z]\.', line):
            # Check if it's a sub-item with nested content
            sub_item_text = line.strip()
            # Preserve bold markers in the text - create paragraph and handle bold
            para = doc.add_paragraph()
            para_format = para.paragraph_format
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(12)
            para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para_format.line_spacing = 1.0
            para_format.left_indent = Inches(0.25)
            para_format.first_line_indent = Inches(-0.25)
            
            # Process text with bold markers
            parts = re.split(r'(\*\*[^*]+\*\*)', sub_item_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                    setup_arial_font(run)
                else:
                    run = para.add_run(part)
                    setup_arial_font(run)
            
            i += 1
            # Check for nested bullet list or continuation
            while i < len(lines):
                next_line = lines[i].strip()
                # Handle nested bullets (indented with spaces)
                if next_line.startswith('   - '):
                    # Nested bullet point (indented)
                    bullet_text = next_line.lstrip('   - ').strip()
                    bullet_para = create_paragraph_with_formatting(doc, f'  • {bullet_text}',
                                                                  indent_left=0.75, first_line_indent=-0.25)
                    i += 1
                elif next_line.startswith('- ') and not next_line.startswith('   -'):
                    # Check context - if next line is another sub-item or section, break
                    if i + 1 < len(lines):
                        peek_line = lines[i + 1].strip()
                        if re.match(r'^[a-z]\.', peek_line) or re.match(r'^##', peek_line) or re.match(r'^###', peek_line):
                            break
                    # Regular bullet - likely nested under sub-item
                    bullet_text = next_line.lstrip('- ').strip()
                    bullet_para = create_paragraph_with_formatting(doc, f'  • {bullet_text}',
                                                                  indent_left=0.75, first_line_indent=-0.25)
                    i += 1
                elif re.match(r'^[a-z]\.', next_line):
                    # Next sub-item
                    break
                elif re.match(r'^##', next_line) or re.match(r'^###', next_line):
                    # Next section
                    break
                elif next_line.startswith('**APPROVAL') or next_line.startswith('**CMDP'):
                    # End of document sections
                    break
                elif not next_line:
                    i += 1
                    continue
                else:
                    # Continuation of current sub-item (regular text)
                    cont_text = next_line
                    cont_para = create_paragraph_with_formatting(doc, cont_text,
                                                               indent_left=0.5, first_line_indent=-0.25)
                    i += 1
            continue
        
        # Handle bullet points
        if line.startswith('- '):
            text = line[2:].strip()
            # Preserve bold markers in bullet text
            para = doc.add_paragraph()
            para_format = para.paragraph_format
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(12)
            para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para_format.line_spacing = 1.0
            para_format.left_indent = Inches(0.5)
            para_format.first_line_indent = Inches(-0.25)
            
            # Process text with bold markers
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            para.add_run('• ')
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                    setup_arial_font(run)
                else:
                    run = para.add_run(part)
                    setup_arial_font(run)
            i += 1
            continue
        
        # Handle APPROVAL section
        if line.startswith('**APPROVAL:**'):
            # Add 4 blank lines before signature per AR 25-50
            for _ in range(4):
                create_paragraph_with_formatting(doc, '', is_empty=True)
            create_paragraph_with_formatting(doc, 'APPROVAL:', bold=True)
            create_paragraph_with_formatting(doc, '', is_empty=True)
            i += 1
            continue
        
        # Handle signature block
        if line.startswith('**[LAST NAME') or line.startswith('[LAST NAME'):
            # Signature block format
            sig_text = line.replace('**', '').replace('[', '').replace(']', '').strip()
            create_paragraph_with_formatting(doc, sig_text, bold=True)
            i += 1
            # Get next two lines for rank and "Commanding"
            if i < len(lines):
                next_line = lines[i].strip().replace('[', '').replace(']', '').strip()
                if next_line:
                    create_paragraph_with_formatting(doc, next_line)
                    i += 1
            if i < len(lines):
                next_line = lines[i].strip()
                if next_line == 'Commanding':
                    create_paragraph_with_formatting(doc, next_line,
                                                   alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                    i += 1
            continue
        
        # Handle CMDP Reference
        if line.startswith('**CMDP Reference:**'):
            ref_text = line.replace('**', '').strip()
            create_paragraph_with_formatting(doc, ref_text, bold=True)
            i += 1
            continue
        
        # Default: regular paragraph
        # Clean up any remaining markdown formatting
        clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
        create_paragraph_with_formatting(doc, clean_line)
        i += 1
    
    # Save document
    output_dir.mkdir(parents=True, exist_ok=True)
    output_filename = Path(md_file_path).stem + '.docx'
    output_path = Path(output_dir) / output_filename
    try:
        doc.save(str(output_path))
        print(f"Converted: {md_file_path.name} -> {output_path.name}")
    except Exception as e:
        print(f"Error saving {output_path}: {e}")
        raise
    return output_path

def main():
    """Main function to convert all SOP markdown files"""
    sops_dir = Path('content/sops')
    output_dir = Path('content/word')
    
    # Ensure output directory exists
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Files to exclude
    exclude_files = {'sops.index.md'}
    
    # Find all .md files
    md_files = [f for f in sops_dir.glob('*.md') 
                if f.name not in exclude_files]
    
    if not md_files:
        print(f"ERROR: No SOP markdown files found in {sops_dir}")
        print(f"Directory exists: {sops_dir.exists()}")
        if sops_dir.exists():
            all_files = list(sops_dir.glob('*.md'))
            print(f"All .md files found: {[f.name for f in all_files]}")
        return
    
    print(f"Found {len(md_files)} SOP files to convert")
    print("Formatting according to AR 25-50 standards:")
    print("  - 12pt Arial font")
    print("  - 1 inch margins on all sides")
    print("  - Single spacing within paragraphs")
    print("  - Double spacing between paragraphs")
    print("  - Proper paragraph alignment and indentation\n")
    
    success_count = 0
    for md_file in sorted(md_files):
        try:
            result = convert_markdown_to_docx(md_file, output_dir)
            if result and result.exists():
                success_count += 1
        except Exception as e:
            print(f"Error converting {md_file.name}: {e}")
            import traceback
            traceback.print_exc()
    
    print(f"\nConversion complete! Successfully converted {success_count}/{len(md_files)} files.")
    print("Note: Image placeholders were skipped. Please add images manually.")

if __name__ == '__main__':
    main()

