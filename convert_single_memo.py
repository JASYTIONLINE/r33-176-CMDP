import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def setup_arial_font(run):
    """Set font to 12pt Arial for a run"""
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    # Set East Asian font to Arial for compatibility
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Arial')

def create_paragraph_with_formatting(doc, text='', bold=False, alignment=None, indent_left=0, first_line_indent=0, is_empty=False):
    """Create a paragraph with AR 25-50 formatting"""
    paragraph = doc.add_paragraph()
    
    if alignment:
        paragraph.alignment = alignment
    
    # Set paragraph formatting per AR 25-50
    # Single-space within paragraphs, double-space between paragraphs
    para_format = paragraph.paragraph_format
    para_format.space_before = Pt(0)
    # Double spacing between paragraphs = 12pt (one line of 12pt font)
    para_format.space_after = Pt(12) if not is_empty else Pt(0)
    para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para_format.line_spacing = 1.0
    
    if indent_left > 0:
        para_format.left_indent = Inches(indent_left)
    if first_line_indent != 0:
        para_format.first_line_indent = Inches(first_line_indent)
    
    if text:
        # Handle bold text markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                setup_arial_font(run)
            else:
                run = paragraph.add_run(part)
                if bold:
                    run.bold = True
                setup_arial_font(run)
    else:
        # Empty paragraph - still set font
        run = paragraph.add_run()
        setup_arial_font(run)
    
    return paragraph

def convert_markdown_to_docx(md_file_path, output_dir):
    """Convert a markdown memo file to DOCX with AR 25-50 formatting"""
    
    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set document margins per AR 25-50: 1 inch on all sides
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set default font to Arial 12pt in Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style.paragraph_format.space_before = Pt(0)
    # AR 25-50: double-space between paragraphs
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.line_spacing = 1.0
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    skip_empty_after_hr = False
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip markdown navigation links
        if line.startswith('[[') and ('|' in line or 'memos.index' in line or 'cmdp-index' in line):
            i += 1
            continue
        
        # Skip horizontal rules (---) - they're visual separators
        if line == '---':
            skip_empty_after_hr = True
            i += 1
            continue
        
        # Skip empty line after horizontal rule
        if skip_empty_after_hr and not line:
            i += 1
            skip_empty_after_hr = False
            continue
        skip_empty_after_hr = False
        
        # Skip image references (user will add these manually)
        if line.startswith('!['):
            i += 1
            continue
        
        # Handle empty lines (no spacing for empty paragraphs)
        if not line:
            create_paragraph_with_formatting(doc, '', is_empty=True)
            i += 1
            continue
        
        # Handle office symbol and date line (right-aligned)
        if 'OFFICE-SYMBOL' in line or (re.match(r'^[A-Z0-9-]+\s+\d{1,2}\s+[A-Z]{3}\s+\d{4}$', line)):
            # Check if it's the office symbol/date format
            if 'OFFICE-SYMBOL' in line or re.search(r'\d{1,2}\s+[A-Z]{3}\s+\d{4}', line):
                create_paragraph_with_formatting(doc, line, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                i += 1
                continue
        
        # Handle MEMORANDUM FOR line - center aligned
        if line.startswith('MEMORANDUM FOR'):
            create_paragraph_with_formatting(doc, line, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue
        
        # Handle SUBJECT line - bold, left aligned
        if line.startswith('SUBJECT:'):
            create_paragraph_with_formatting(doc, line, bold=True)
            i += 1
            continue
        
        # Handle DISTRIBUTION line
        if line.startswith('**DISTRIBUTION:**') or line.startswith('DISTRIBUTION:'):
            dist_text = line.replace('**', '')
            create_paragraph_with_formatting(doc, dist_text, bold=True)
            i += 1
            continue
        
        # Handle CMDP Reference
        if line.startswith('**CMDP Reference:**'):
            ref_text = line.replace('**', '')
            create_paragraph_with_formatting(doc, ref_text, bold=True)
            i += 1
            continue
        
        # Handle numbered paragraphs (1., 2., etc.)
        if re.match(r'^\d+\.', line):
            # Check if next line is a sub-item
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if re.match(r'^\s+[a-z]\.', next_line):
                    # Main numbered paragraph
                    create_paragraph_with_formatting(doc, line)
                    # Process sub-items
                    i += 1
                    while i < len(lines) and re.match(r'^\s+[a-z]\.', lines[i].strip()):
                        sub_line = lines[i].strip()
                        # Remove bold markers but keep content
                        sub_line_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', sub_line)
                        para = create_paragraph_with_formatting(doc, sub_line_clean, indent_left=0.5)
                        i += 1
                    continue
            # Regular numbered paragraph
            create_paragraph_with_formatting(doc, line)
            i += 1
            continue
        
        # Handle bullet points (duties section)
        if line.startswith('- '):
            text = line[2:].strip()
            # Remove bold markers but preserve structure
            text_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            para = create_paragraph_with_formatting(doc, text_clean, indent_left=0.5, first_line_indent=-0.25)
            # Add bullet character
            if para.runs:
                para.runs[0].text = '• ' + para.runs[0].text
            i += 1
            continue
        
        # Handle signature block - AR 25-50 requires 4 blank lines before signature
        if line == 'Respectfully,':
            # Add 4 blank lines before signature per AR 25-50
            for _ in range(4):
                create_paragraph_with_formatting(doc, '', is_empty=True)
            create_paragraph_with_formatting(doc, line)
            create_paragraph_with_formatting(doc, '', is_empty=True)
            i += 1
            continue
        
        # Handle "Commanding" line - right aligned
        if line == 'Commanding':
            create_paragraph_with_formatting(doc, line, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            i += 1
            continue
        
        # Default: regular paragraph
        # Clean up any remaining markdown formatting
        clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
        create_paragraph_with_formatting(doc, clean_line)
        i += 1
    
    # Save document
    output_filename = Path(md_file_path).stem + '.docx'
    output_path = Path(output_dir) / output_filename
    doc.save(str(output_path))
    print(f"Converted: {md_file_path.name} -> {output_path.name}")
    return output_path

if __name__ == '__main__':
    # Convert just the 10.1.1 memo
    md_file = Path('content/memos/10.1.1-maint-officer.md')
    output_dir = Path('content/word')
    
    print("Converting 10.1.1-maint-officer.md to Word format...")
    print("Formatting according to AR 25-50 standards:")
    print("  - 12pt Arial font")
    print("  - 1 inch margins on all sides")
    print("  - Single spacing within paragraphs")
    print("  - Double spacing between paragraphs")
    print("  - Proper paragraph alignment\n")
    
    try:
        convert_markdown_to_docx(md_file, output_dir)
        print("\nConversion complete! File ready for QC.")
        print("Note: Image placeholders were skipped. Please add images manually.")
    except Exception as e:
        print(f"Error converting file: {e}")
        import traceback
        traceback.print_exc()

